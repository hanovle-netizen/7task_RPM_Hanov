from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Поля страницы
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# Стили шрифта
def set_font(run, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=14 if level == 1 else 13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def para(text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
    r2 = p.add_run(text)
    set_font(r2)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 180)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        # Серый фон
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F0F0')
        pPr.append(shd)

# ══════════════════════════════════════════
#  ТИТУЛ
# ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Лабораторная работа №7')
set_font(r, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Поведенческие паттерны проектирования:\nНаблюдатель, Стратегия, Шаблонный метод')
set_font(r, size=14, bold=True)

doc.add_paragraph()

# ══════════════════════════════════════════
#  РАЗДЕЛ 1 — КАК РАБОТАЕТ КОД
# ══════════════════════════════════════════
heading('1. Как работает код — объяснение по файлам')

heading('1.1. MetricData.cs — данные события', level=2)
para(
    'Простой класс-контейнер, который хранит информацию о метрике в момент события: '
    'название (CPU_Load, Memory_Usage...), зафиксированное значение, порог и время.',
    indent=True
)
code_block('''\
public class MetricData(string metricName, double value, double threshold, DateTime timestamp)
{
    public string MetricName { get; } = metricName;
    public double Value      { get; } = value;
    public double Threshold  { get; } = threshold;
    public DateTime Timestamp{ get; } = timestamp;
}''')

heading('1.2. EventMonitor.cs — издатель (паттерн Наблюдатель)', level=2)
para(
    'Это «сердце» системы. Класс проверяет значение метрики: если оно превышает порог — '
    'публикует событие OnMetricExceeded. Все подписчики получают уведомление автоматически.',
    indent=True
)
para(
    'Ключевые строки: ',
)
code_block('''\
public event MetricEventHandler? OnMetricExceeded;  // объявляем событие

public void CheckMetric(string metricName, double value, double threshold)
{
    if (value > threshold)
    {
        var eventData = new MetricData(metricName, value, threshold, DateTime.Now);
        OnMetricExceeded?.Invoke(                          // публикуем событие
            new MetricEventArgs(metricName + "_Exceeded", eventData));
    }
}''')
para('?.Invoke — безопасный вызов: если подписчиков нет, исключения не будет.', indent=True)

heading('1.3. IFormatStrategy / TextFormat / JsonFormat / HtmlFormat — паттерн Стратегия', level=2)
para(
    'Интерфейс IFormatStrategy задаёт один метод Format(). '
    'Три конкретных класса реализуют его по-разному: '
    'текст, JSON и HTML. Можно добавить XML или любой другой формат, '
    'не трогая остальной код.',
    indent=True
)
code_block('''\
public interface IFormatStrategy
{
    string Format(string message, DateTime timestamp);
}

// Текст:  [2026-03-31 14:00:00] сообщение
// JSON:   {"timestamp":"...","message":"..."}
// HTML:   <div class="alert"><span>...</span></div>''')

heading('1.4. EventHandlerBase.cs — паттерн Шаблонный метод + контекст Стратегии', level=2)
para(
    'Абстрактный базовый класс. Содержит шаблонный метод ProcessEvent, '
    'который фиксирует алгоритм из трёх шагов. '
    'Также хранит текущую стратегию форматирования.',
    indent=True
)
code_block('''\
public void ProcessEvent(MetricEventArgs e)   // <-- шаблонный метод
{
    var message = FormatMessage(e.EventType, e.Data);  // шаг 1: форматировать
    SendMessage(message);                               // шаг 2: отправить
    LogResult();                                        // шаг 3: залогировать (hook)
}

protected abstract string FormatMessage(string type, object data);  // подкласс реализует
protected abstract void   SendMessage(string message);               // подкласс реализует
protected virtual  void   LogResult() { }                            // hook — необязателен''')

heading('1.5. ConsoleHandler / FileHandler / EmailHandler — конкретные обработчики', level=2)
para(
    'Каждый класс наследует EventHandlerBase и реализует только детали: '
    'куда именно отправить сообщение. Алгоритм (порядок шагов) один и тот же для всех.',
    indent=True
)
code_block('''\
// ConsoleHandler — выводит в консоль жёлтым цветом
protected override void SendMessage(string message)
{
    Console.ForegroundColor = ConsoleColor.Yellow;
    Console.WriteLine($"[ConsoleHandler:{_name}] {message}");
    Console.ResetColor();
}

// FileHandler — дописывает строку в файл alerts.html
protected override void SendMessage(string message)
{
    File.AppendAllText(_filePath, message + Environment.NewLine);
}''')

heading('1.6. Program.cs — демонстрация работы всей системы', level=2)
para(
    'Точка входа. Создаёт монитор, обработчики с разными стратегиями, '
    'подписывает их на событие и симулирует метрики.',
    indent=True
)
code_block('''\
var monitor = new EventMonitor();

// Создаём обработчики с разными стратегиями форматирования
var consoleText = new ConsoleHandler("TEXT", new TextFormatStrategy());
var consoleJson = new ConsoleHandler("JSON", new JsonFormatStrategy());
var fileHtml    = new FileHandler("alerts.html", new HtmlFormatStrategy());

// Подписка (Attach) — паттерн Наблюдатель
monitor.OnMetricExceeded += consoleText.ProcessEvent;
monitor.OnMetricExceeded += consoleJson.ProcessEvent;
monitor.OnMetricExceeded += fileHtml.ProcessEvent;

// Симуляция: 45 < 80 — норма, 95 > 80 — событие!
monitor.CheckMetric("CPU_Load", 45.0, 80.0);
monitor.CheckMetric("CPU_Load", 95.0, 80.0);

// Смена стратегии в runtime
consoleJson.SetStrategy(new HtmlFormatStrategy());

// Отписка (Detach)
monitor.OnMetricExceeded -= consoleText.ProcessEvent;''')

# ══════════════════════════════════════════
#  РАЗДЕЛ 2 — КОНТРОЛЬНЫЕ ВОПРОСЫ
# ══════════════════════════════════════════
heading('2. Ответы на контрольные вопросы')

questions = [
    (
        'Вопрос 1. Какой паттерн использован для доставки уведомлений? Почему он оптимален?',
        '''Паттерн «Наблюдатель» (Observer). Он оптимален по нескольким причинам:

• EventMonitor (издатель) ничего не знает о конкретных типах подписчиков — только о том, что у них есть метод ProcessEvent с нужной сигнатурой. Это обеспечивает слабую связанность.
• Число подписчиков можно менять динамически в runtime через += и -=, не трогая код издателя.
• В C# конструкция event автоматически защищает от состояния гонки: компилятор делает снимок списка подписчиков перед вызовом, поэтому одновременная подписка/отписка не вызовет исключение.'''
    ),
    (
        'Вопрос 2. Как «Стратегия» обеспечивает гибкость форматирования?',
        '''Интерфейс IFormatStrategy скрывает детали форматирования за единым методом Format(). EventHandlerBase работает только с этим интерфейсом и не знает, какой именно формат используется.

Преимущество смены стратегии в runtime: если ночью нужен краткий текстовый лог, а днём — подробный JSON для системы мониторинга, можно вызвать SetStrategy() прямо во время работы приложения без перезапуска и без изменения кода обработчика.

В коде это выглядит так:
    consoleJson.SetStrategy(new HtmlFormatStrategy());
После этой строки тот же обработчик начинает выдавать HTML вместо JSON.'''
    ),
    (
        'Вопрос 3. Роль шаблонного метода в EventHandlerBase? Что фиксировано, что переопределяется?',
        '''Шаблонный метод ProcessEvent фиксирует порядок трёх шагов алгоритма обработки события:
    1. FormatMessage — сформировать текст сообщения
    2. SendMessage   — отправить в канал доставки
    3. LogResult     — залогировать результат

Фиксировано: сама последовательность вызовов (1→2→3) — она одинакова для всех обработчиков.

Переопределяется в подклассах:
• FormatMessage и SendMessage — абстрактные методы, ОБЯЗАНЫ быть реализованы.
• LogResult — hook-метод (виртуальный), по умолчанию пустой. Подкласс МОЖЕТ переопределить его для дополнительного поведения (ConsoleHandler и FileHandler переопределяют — выводят подтверждение).'''
    ),
    (
        'Вопрос 4. Как три паттерна взаимодействуют? Поток данных от события до обработки.',
        '''Поток данных при вызове monitor.CheckMetric("CPU_Load", 95.0, 80.0):

1. EventMonitor.CheckMetric видит превышение (95 > 80)
2. Создаёт MetricData и MetricEventArgs с данными
3. Вызывает OnMetricExceeded?.Invoke(args)  ← НАБЛЮДАТЕЛЬ публикует событие
4. Каждый подписчик получает вызов своего ProcessEvent(args)
5. ProcessEvent — шаблонный метод (ШАБЛОННЫЙ МЕТОД):
   а) вызывает FormatMessage → тот обращается к _formatStrategy.Format()  ← СТРАТЕГИЯ
   б) вызывает SendMessage → ConsoleHandler пишет в консоль, FileHandler в файл
   в) вызывает LogResult → выводит подтверждение

Итог: Observer доставляет событие, Template Method управляет алгоритмом, Strategy форматирует сообщение.'''
    ),
    (
        'Вопрос 5. Какие принципы SOLID соблюдаются?',
        '''SRP (Single Responsibility — единственная ответственность):
• EventMonitor отвечает только за мониторинг метрик
• ConsoleHandler только за вывод в консоль
• TextFormatStrategy только за текстовое форматирование

OCP (Open/Closed — открыт для расширения, закрыт для изменения):
• Новый формат (XML) = новый класс XmlFormatStrategy, существующий код не меняется
• Новый канал (SMS) = новый класс SmsHandler, существующий код не меняется

LSP (Liskov Substitution — подстановка Лисков):
• Любая стратегия взаимозаменяема через IFormatStrategy
• Любой обработчик работает одинаково через EventHandlerBase

DIP (Dependency Inversion — инверсия зависимостей):
• EventHandlerBase зависит от абстракции IFormatStrategy, а не от конкретных классов'''
    ),
    (
        'Вопрос 6. Как расширить систему (новый обработчик SMS, новый формат XML)?',
        '''Добавить SMS-обработчик — 3 шага, без изменения существующего кода:

    public class SmsHandler(string phone, IFormatStrategy strategy)
        : EventHandlerBase(strategy)
    {
        protected override string FormatMessage(string type, object data) =>
            _formatStrategy.Format($"[{type}] {data}", DateTime.Now);

        protected override void SendMessage(string message) =>
            Console.WriteLine($"[SMS -> {phone}]: {message}"); // здесь реальный SMS API
    }

Добавить XML-формат — 1 класс:

    public class XmlFormatStrategy : IFormatStrategy
    {
        public string Format(string message, DateTime timestamp) =>
            $"<alert><time>{timestamp:O}</time><msg>{message}</msg></alert>";
    }

Затем в Program.cs:
    monitor.OnMetricExceeded += new SmsHandler("+7999...", new XmlFormatStrategy()).ProcessEvent;'''
    ),
    (
        'Вопрос 7. Чем «Стратегия» отличается от «Шаблонного метода»?',
        '''Ключевое различие — механизм расширения:

«Шаблонный метод» использует НАСЛЕДОВАНИЕ:
• Алгоритм определён в базовом классе
• Подкласс переопределяет отдельные шаги через override
• Структура фиксирована на этапе компиляции
• Подходит, когда общий алгоритм стабилен, меняются только детали

«Стратегия» использует КОМПОЗИЦИЮ:
• Алгоритм вынесен в отдельный объект-стратегию
• Стратегию можно заменить в runtime одним вызовом SetStrategy()
• Более гибкий подход
• Подходит, когда нужно переключать поведение во время работы программы

В нашей системе оба паттерна работают вместе: Template Method фиксирует порядок шагов (форматировать → отправить → залогировать), а Strategy гибко меняет реализацию шага форматирования.'''
    ),
    (
        'Вопрос 8. Проблемы Observer в многопоточной среде и их решения.',
        '''А. Состояние гонки (Race Condition):
• Проблема: поток A перебирает список подписчиков, поток B одновременно добавляет/удаляет подписчика → InvalidOperationException
• Решение: конструкция event в C# автоматически делает снимок делегата перед вызовом — в нашем коде это уже защищено

Б. Блокировка издателя (Latency):
• Проблема: FileHandler или EmailHandler выполняет долгую операцию (диск, сеть) → EventMonitor заблокирован, остальные подписчики ждут
• Решение: вызывать обработчики асинхронно через Task.Run(() => handler.ProcessEvent(args))

В. Взаимная блокировка (Deadlock):
• Проблема: издатель держит lock на списке подписчиков, подписчик внутри Update пытается отписаться — тоже требует того же lock → зависание
• Решение: использовать ConcurrentBag<T> или CopyOnWrite-коллекции

Г. Порядок уведомлений:
• Проблема: в многопоточной среде нет гарантии порядка оповещения
• Решение: если порядок важен — использовать очередь событий (Queue) с одним потоком-обработчиком'''
    ),
]

for q_title, q_answer in questions:
    heading(q_title, level=2)
    for line in q_answer.strip().split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        if line.startswith('    '):
            # код внутри ответа
            run = p.add_run(line.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 30, 180)
        else:
            run = p.add_run(line)
            set_font(run)

output = r"C:\Users\Zola\Desktop\Поведенческие паттерны проектирования. Наблюдатель, Стратегия, Шаблонный метод\MonitoringSystem\Отчёт_ЛР7.docx"
doc.save(output)
print(f"Файл сохранён: {output}")
